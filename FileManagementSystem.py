import tkinter as tk
from tkinter import simpledialog, messagebox, ttk, scrolledtext
import os
import time
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import Login
import Encrypt
from docx import Document
import shutil
import Database
import psutil
import Encrypt_rsa
import matplotlib.pyplot as plt


class LoginWindow:
    def __init__(self, root):
        self.root = root
        self.root.title("Log in")
        self.root.geometry("400x300")

        # tittle
        tk.Label(self.root, text="Log in to the file management system", font=("Arial", 14)).pack(pady=10)

        # input user Name
        tk.Label(self.root, text="User Name：").pack(pady=5)
        self.username_entry = tk.Entry(self.root)
        self.username_entry.pack()

        # password input
        tk.Label(self.root, text="Password：").pack(pady=5)
        self.password_entry = tk.Entry(self.root, show="*")
        self.password_entry.pack()

        # button
        tk.Button(self.root, text="Log in", command=self.login).pack(pady=10)
        tk.Button(self.root, text="Sign Up", command=self.open_register_window).pack(pady=5)

    def login(self):
        """user login"""
        username = self.username_entry.get().strip()
        password = self.password_entry.get().strip()

        if not username or not password:
            messagebox.showerror("login failure", "The user name and password cannot be empty！")
            return

        if Login.login_user(username, password):
            messagebox.showinfo("login successfully", f"welcome {username}！")
            self.root.destroy()
            MainWindow(username)
        else:
            messagebox.showerror("login failure", "The user name or password is incorrect. Please try again.")

    def open_register_window(self):
        """Open the registration window"""
        self.root.withdraw()  # Hide current window
        RegisterWindow(self.root)  # Open the registration window




class RegisterWindow:
    def __init__(self, parent):
        self.parent = parent
        self.root = tk.Toplevel()
        self.root.title("sign up")
        self.root.geometry("400x300")

        tk.Label(self.root, text="Register a new user", font=("Arial", 14)).pack(pady=10)

        tk.Label(self.root, text="user name：").pack(pady=5)
        self.username_entry = tk.Entry(self.root)
        self.username_entry.pack()
        self.username_entry.bind("<KeyRelease>", self.check_username)

        self.username_status = tk.Label(self.root, text="", fg="red")
        self.username_status.pack()

        tk.Label(self.root, text="password：").pack(pady=5)
        self.password_entry = tk.Entry(self.root, show="*")
        self.password_entry.pack()
        self.password_entry.config(state=tk.DISABLED)

        tk.Button(self.root, text="sign up", command=self.register).pack(pady=10)
        tk.Button(self.root, text="return", command=self.go_back).pack(pady=5)

    def check_username(self, event):
        """Check whether the user name exists in real time"""
        username = self.username_entry.get().strip()
        if not username:
            self.username_status.config(text="The user name cannot be empty", fg="red")
            self.password_entry.config(state=tk.DISABLED)
            return

        if Login.check_user_exists(username):
            self.username_status.config(text="The user name already exists. Please change it", fg="red")
            self.password_entry.config(state=tk.DISABLED)
        else:
            self.username_status.config(text="User name available", fg="green")
            self.password_entry.config(state=tk.NORMAL)

    def register(self):
        """user register"""
        username = self.username_entry.get().strip()
        password = self.password_entry.get().strip()

        if not username or not password:
            messagebox.showerror("fail to register", "The user name and password cannot be empty！")
            return

        if Login.register_user(username, password):
            messagebox.showinfo("registered successfully", f"user {username} Registration successful, please login！")
            self.go_back()
        else:
            messagebox.showerror("fail to register", f"username {username} Already exists. Please select another user name.")

    def go_back(self):
        """Return to login window"""
        self.root.destroy()
        self.parent.deiconify()


class MainWindow:
    def __init__(self, username):
        self.root = tk.Tk()
        self.root.title(f"file management system - {username}")
        self.root.geometry("1200x800")
        self.username = username
        self.check_keys()

        # Main interface layout
        self.create_layout()

        self.selected_file_path = None

        self.backup_directory = os.path.join(os.getcwd(), "backups")
        if not os.path.exists(self.backup_directory):
            os.makedirs(self.backup_directory)

        self.populate_file_tree()

        self.update_disk_info()

        self.root.mainloop()

    def check_keys(self):
            """Check whether the public key and private key file exist. If no, the file is generated"""
            if not os.path.exists("private_key.pem") or not os.path.exists("public_key.pem"):
                Encrypt_rsa.generate_rsa_keys()
                messagebox.showinfo("Keys Generated", "RSA keys have been generated.")

    def create_layout(self):
        """Create the layout of the main screen"""
        # Ministry search bar
        self.search_frame = tk.Frame(self.root)
        self.search_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=5)
        tk.Label(self.search_frame, text="Search for files or folders：").pack(side=tk.LEFT)
        self.search_entry = tk.Entry(self.search_frame, width=40)
        self.search_entry.pack(side=tk.LEFT, padx=5)
        tk.Button(self.search_frame, text="search", command=self.search_file).pack(side=tk.LEFT, padx=5)

        # Left: File tree view + scroll bar
        self.file_tree_frame = tk.Frame(self.root)
        self.file_tree_frame.pack(side=tk.LEFT, fill=tk.BOTH, padx=10, pady=10)

        self.file_tree_scrollbar_y = tk.Scrollbar(self.file_tree_frame, orient=tk.VERTICAL)
        self.file_tree_scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)

        self.file_tree_scrollbar_x = tk.Scrollbar(self.file_tree_frame, orient=tk.HORIZONTAL)
        self.file_tree_scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)

        self.file_tree = ttk.Treeview(
            self.file_tree_frame,
            yscrollcommand=self.file_tree_scrollbar_y.set,
            xscrollcommand=self.file_tree_scrollbar_x.set
        )
        self.file_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.file_tree_scrollbar_y.config(command=self.file_tree.yview)
        self.file_tree_scrollbar_x.config(command=self.file_tree.xview)

        self.file_tree.bind("<<TreeviewSelect>>", self.on_file_select)
        self.file_tree.bind("<<TreeviewOpen>>", self.on_folder_expand)

        # Right: Function buttons and log area
        self.right_frame = tk.Frame(self.root)
        self.right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        # function button
        self.button_frame = tk.Frame(self.right_frame)
        self.button_frame.pack(side=tk.TOP, fill=tk.X, pady=10)
        # Add a time record label box
        self.time_log_frame = tk.LabelFrame(self.right_frame, text="Time Log", padx=10, pady=10)
        self.time_log_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=10)

        # Initializes the time record label
        self.time_labels = {
            "AES Encrypt Time": tk.Label(self.time_log_frame, text="AES Encrypt Time: Not Recorded"),
            "AES Decrypt Time": tk.Label(self.time_log_frame, text="AES Decrypt Time: Not Recorded"),
            "RSA Encrypt Time": tk.Label(self.time_log_frame, text="RSA Encrypt Time: Not Recorded"),
            "RSA Decrypt Time": tk.Label(self.time_log_frame, text="RSA Decrypt Time: Not Recorded")
        }

        # Place the labels in order
        for label in self.time_labels.values():
            label.pack(anchor="w", pady=2)

        # First row button
        tk.Button(self.button_frame, text="Refresh", width=15, command=self.refresh_file_tree).grid(row=0, column=0,
                                                                                                    padx=5, pady=5)
        tk.Button(self.button_frame, text="Encrypted File", width=15, command=self.encrypt_file).grid(row=0, column=1,
                                                                                                      padx=5, pady=5)
        tk.Button(self.button_frame, text="Decrypted file", width=15, command=self.decrypt_file).grid(row=0,
                                                                                                      column=2,
                                                                                                      padx=5, pady=5)
        tk.Button(self.button_frame, text="Delete File", width=15, command=self.delete_file).grid(row=0, column=3,
                                                                                                  padx=5, pady=5)

        # Second row button
        tk.Button(self.button_frame, text="Clear Log", width=15, command=self.clear_log).grid(row=1, column=0, padx=5,
                                                                                              pady=5)
        tk.Button(self.button_frame, text="Rename File", width=15, command=self.rename_file).grid(row=1, column=1,
                                                                                                  padx=5, pady=5)
        tk.Button(self.button_frame, text="Read File", width=15, command=self.read_write_file).grid(row=1, column=2,
                                                                                                    padx=5, pady=5)

        # Third row button
        tk.Button(self.button_frame, text="Backup File", width=15, command=self.backup_file).grid(row=2, column=0,
                                                                                                  padx=5, pady=5)
        tk.Button(self.button_frame, text="Restore File", width=15, command=self.restore_file).grid(row=2, column=1,
                                                                                                    padx=5, pady=5)
        tk.Button(self.button_frame, text="Exit", width=15, command=self.root.quit).grid(row=2, column=2, padx=5,

                                                                                         pady=5)
        tk.Button(self.button_frame, text="RSA Encrypt File", width=15, command=self.rsa_encrypt_file).grid(row=3,
                                                                                                            column=0,
                                                                                                            padx=5,
                                                                                                            pady=5)
        tk.Button(self.button_frame, text="RSA Decrypt File", width=15, command=self.rsa_decrypt_file).grid(row=3,
                                                                                                            column=1,
                                                                                                            padx=5,
                                                                                                            pady=5)
        tk.Button(self.button_frame, text="Compare AES vs RSA", width=15, command=self.compare_aes_rsa).grid(row=3,
                                                                                                             column=2,
                                                                                                             padx=5,
                                                                                                             pady=5)

        # Display disk capacity
        self.disk_info_frame = tk.LabelFrame(self.right_frame, text="Disk Space Information", padx=10, pady=10)
        self.disk_info_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=10)

        self.c_drive_label = tk.Label(self.disk_info_frame, text="C Drive: ")
        self.c_drive_label.pack(anchor="w")

        self.c_drive_bar = ttk.Progressbar(self.disk_info_frame, orient="horizontal", length=300, mode="determinate")
        self.c_drive_bar.pack(fill=tk.X, padx=10, pady=5)

        self.d_drive_label = tk.Label(self.disk_info_frame, text="D Drive: ")
        self.d_drive_label.pack(anchor="w")

        self.d_drive_bar = ttk.Progressbar(self.disk_info_frame, orient="horizontal", length=300, mode="determinate")
        self.d_drive_bar.pack(fill=tk.X, padx=10, pady=5)

        # File information display area
        self.file_info_frame = tk.LabelFrame(self.right_frame, text="file information", padx=10, pady=10)
        self.file_info_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=10)

        self.file_info_labels = {
            "File Name": tk.Label(self.file_info_frame, text="File Name: "),
            "File Path": tk.Label(self.file_info_frame, text="File Path: "),
            "File Type": tk.Label(self.file_info_frame, text="File Type: "),
            "File Size": tk.Label(self.file_info_frame, text="File Size: "),
            "Last Modified Time": tk.Label(self.file_info_frame, text="Last Modified Time: "),
        }
        for label in self.file_info_labels.values():
            label.pack(anchor="w")

        self.log_frame = tk.Frame(self.right_frame)
        self.log_frame.pack(side=tk.BOTTOM, fill=tk.BOTH, expand=True, pady=10)

        self.log_scrollbar = tk.Scrollbar(self.log_frame, orient=tk.VERTICAL)
        self.log_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.log_text = tk.Text(self.log_frame, height=15, wrap=tk.WORD, yscrollcommand=self.log_scrollbar.set)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.log_scrollbar.config(command=self.log_text.yview)
        self.log_text.config(state=tk.DISABLED)

    def populate_file_tree(self):
        """Initializes the file tree to display the local hard disk and folders"""
        drives = [f"{chr(d)}:/" for d in range(65, 91) if os.path.exists(f"{chr(d)}:/")]
        for drive in drives:
            drive_node = self.file_tree.insert("", "end", text=drive, values=[drive])
            self.add_placeholder(drive_node)

    def update_disk_info(self):
        """Example Update the disk capacity information of drive C and drive D"""
        # Obtain disk usage information about disk C and disk D
        c_drive_usage = psutil.disk_usage("C:/")
        d_drive_usage = psutil.disk_usage("D:/")

        #Formatted capacity information
        c_drive_used = c_drive_usage.used / (1024**3)
        c_drive_total = c_drive_usage.total / (1024**3)
        c_drive_percent = (c_drive_used / c_drive_total) * 100

        d_drive_used = d_drive_usage.used / (1024**3)
        d_drive_total = d_drive_usage.total / (1024**3)
        d_drive_percent = (d_drive_used / d_drive_total) * 100


        self.c_drive_bar["value"] = c_drive_percent
        self.d_drive_bar["value"] = d_drive_percent

        self.c_drive_label.config(text=f"C Drive: {c_drive_used:.2f} GB used / {c_drive_total:.2f} GB total")
        self.d_drive_label.config(text=f"D Drive: {d_drive_used:.2f} GB used / {d_drive_total:.2f} GB total")


    def add_placeholder(self, node):
        """Add a placeholder for the directory node"""
        if not self.file_tree.get_children(node):
            self.file_tree.insert(node, "end", text="...", values=["placeholder"])

    def populate_folder(self, parent_node, folder_path):
        """Load folder contents"""
        self.file_tree.delete(*self.file_tree.get_children(parent_node))
        try:
            for item in os.listdir(folder_path):
                item_path = os.path.join(folder_path, item)
                if os.path.isdir(item_path):
                    folder_node = self.file_tree.insert(parent_node, "end", text=item, values=[item_path])
                    self.add_placeholder(folder_node)
                elif os.path.isfile(item_path):
                    self.file_tree.insert(parent_node, "end", text=item, values=[item_path])
        except (PermissionError, FileNotFoundError):
            self.log_message(f"Unable to access folder：{folder_path}")

    def on_folder_expand(self, event):
        """Folder expansion event"""
        selected_item = self.file_tree.selection()
        if selected_item:
            selected_item = selected_item[0]
            folder_path = self.file_tree.item(selected_item, "values")[0]
            if folder_path and folder_path != "placeholder":
                self.populate_folder(selected_item, folder_path)




    def refresh_file_tree(self):
        """Refresh file tree"""
        self.file_tree.delete(*self.file_tree.get_children())
        self.populate_file_tree()
        self.log_message("The file tree is refreshed")

    def on_file_select(self, event):
        """Handle file selection events in the file tree"""
        selected_item = self.file_tree.selection()
        if not selected_item:
            self.log_message("No files or folders are selected")
            return

        selected_item = selected_item[0]
        values = self.file_tree.item(selected_item, "values")

        if not values:
            self.log_message("No valid file or folder is selected")
            self.selected_file_path = None
            return

        file_path = values[0]
        if file_path == "placeholder":
            self.log_message("This is a placeholder item and cannot be selected")
            return

        if os.path.isfile(file_path):
            self.selected_file_path = file_path
            self.log_message(f"file selected：{file_path}")
            self.display_file_info(file_path)
        elif os.path.isdir(file_path):
            self.selected_file_path = None
            self.log_message(f"Select folder：{file_path}")
            self.clear_file_info()

    def backup_file(self):
        """Back up selected files"""
        if not self.selected_file_path:
            messagebox.showerror("error", "Please select a file first！")
            return

        file_name = os.path.basename(self.selected_file_path)
        backup_path = os.path.join(self.backup_directory, f"{file_name}.backup_{int(time.time())}")

        try:
            shutil.copy2(self.selected_file_path, backup_path)
            self.log_message(f"The file has been backed up to：{backup_path}")
            messagebox.showinfo("Backup successfully created", f"The file has been backed up to：{backup_path}")
        except Exception as e:
            self.log_message(f"Backup failed：{e}")
            messagebox.showerror("Backup failed", f"Backup failed：{e}")

    def restore_file(self):
        """Restore selected file """
        if not self.selected_file_path:
            messagebox.showerror("error", "Please select a file first！")
            return

        file_name = os.path.basename(self.selected_file_path)
        backups = [f for f in os.listdir(self.backup_directory) if f.startswith(file_name)]

        if not backups:
            messagebox.showerror("error", "No backup file found！")
            return

        backups.sort(reverse=True)
        latest_backup = os.path.join(self.backup_directory, backups[0])

        # Confirm whether the current file is overwritten
        confirm = messagebox.askyesno("overwrite without confirmation", f"Make sure you want to overwrite the current file with a backup：\n{self.selected_file_path}？")
        if confirm:  # If the user chooses "Yes"
            try:
                shutil.copy2(latest_backup, self.selected_file_path)
                self.log_message(f"The file was restored to：{self.selected_file_path}")
                messagebox.showinfo("Restore successful", f"The file was restored to：{self.selected_file_path}")
            except Exception as e:
                messagebox.showerror("REVERT", f"Unrecoverable file：{str(e)}")
                self.log_message(f"REVERT：{str(e)}")
        else:
            self.log_message("Unrestored operation")

    def display_file_info(self, file_path):
        """Display file information """
        file_name = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        file_type = os.path.splitext(file_name)[1] or "Unknown"
        last_modified = time.ctime(os.path.getmtime(file_path))

        self.file_info_labels["File Name"].config(text=f"File Name: {file_name}")
        self.file_info_labels["File Path"].config(text=f"File Path: {file_path}")
        self.file_info_labels["File Type"].config(text=f"File Type: {file_type}")
        self.file_info_labels["File Size"].config(text=f"File Size: {file_size} bytes")
        self.file_info_labels["Last Modified Time"].config(text=f"Last Modified Time: {last_modified}")

    def clear_file_info(self):
        """Clear the file information display """
        for key in self.file_info_labels:
            self.file_info_labels[key].config(text=f"{key}: ")

    def log_message(self, message):
        """Displays messages in the log area """
        self.log_text.config(state=tk.NORMAL)
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.config(state=tk.DISABLED)
    def search_file(self):
        """Search for files or folders """
        keyword = self.search_entry.get().strip().lower()
        if not keyword:
            messagebox.showwarning("Search warning", "Please enter the search keyword")
            return
        matched_items = self.search_in_tree(self.file_tree.get_children(), keyword)
        if matched_items:
            for item in matched_items:
                self.file_tree.see(item)
                self.file_tree.selection_set(item)
                self.log_message(f"match：{self.file_tree.item(item, 'text')}")
        else:
            self.log_message("No matching file or folder was found")

    def search_in_tree(self, items, keyword):
        """Search recursively for matches in the file tree """
        matched_items = []
        for item in items:
            if keyword in self.file_tree.item(item, "text").lower():
                matched_items.append(item)
            matched_items.extend(self.search_in_tree(self.file_tree.get_children(item), keyword))
        return matched_items

    def log_message(self, message):
        """Displays messages in the log area """
        self.log_text.config(state=tk.NORMAL)
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.config(state=tk.DISABLED)

    def clear_log(self):
        """Clear the log output area """
        self.log_text.config(state=tk.NORMAL)
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state=tk.DISABLED)

    def rsa_encrypt_file(self):
        """Use RSA to encrypt files and record the time"""
        if not self.selected_file_path:
            messagebox.showerror("Error", "Please select a file first")
            return

        try:
            public_key = Encrypt_rsa.load_public_key()
        except FileNotFoundError:
            messagebox.showerror("Error", "Public key not found. Please generate RSA keys first.")
            return

        start_time = time.time()
        encrypted_file = self.selected_file_path + ".rsa.lock"
        Encrypt_rsa.encrypt_file_rsa(self.selected_file_path, encrypted_file, public_key)
        elapsed_time = time.time() - start_time

        self.time_labels["RSA Encrypt Time"].config(text=f"RSA Encrypt Time: {elapsed_time:.4f} seconds")
        self.log_message(f"File RSA-encrypted in {elapsed_time:.4f} seconds: {encrypted_file}")

    def rsa_decrypt_file(self):
        """Use RSA to decrypt the file and record the time """
        if not self.selected_file_path:
            messagebox.showerror("Error", "Please select a file first")
            return

        try:
            private_key = Encrypt_rsa.load_private_key()
        except FileNotFoundError:
            messagebox.showerror("Error", "Private key not found. Please generate RSA keys first.")
            return

        start_time = time.time()
        decrypted_file = self.selected_file_path.replace(".rsa.lock", ".decrypted")
        Encrypt_rsa.decrypt_file_rsa(self.selected_file_path, decrypted_file, private_key)
        elapsed_time = time.time() - start_time

        # Update time record
        self.time_labels["RSA Decrypt Time"].config(text=f"RSA Decrypt Time: {elapsed_time:.4f} seconds")
        self.log_message(f"File RSA-decrypted in {elapsed_time:.4f} seconds: {decrypted_file}")

    def encrypt_file(self):
        """Encrypt Select the file and record the time"""
        if not self.selected_file_path:
            messagebox.showerror("Error", "Please select a file first")
            return

        password = simpledialog.askstring("Encrypt File", "Enter the encryption password:", show="*")
        if password:
            start_time = time.time()
            encrypted_file = self.selected_file_path + ".lock"
            Encrypt.encrypt_file(self.selected_file_path, encrypted_file, password)
            elapsed_time = time.time() - start_time


            self.time_labels["AES Encrypt Time"].config(text=f"AES Encrypt Time: {elapsed_time:.4f} seconds")
            self.log_message(f"File encrypted in {elapsed_time:.4f} seconds: {encrypted_file}")

    def decrypt_file(self):
        """Decrypt the selected file and record the time"""
        if not self.selected_file_path or not self.selected_file_path.endswith(".lock"):
            messagebox.showerror("Error", "Please select an encrypted file")
            return

        password = simpledialog.askstring("Decrypt File", "Enter the decryption password:", show="*")
        if password:
            start_time = time.time()
            decrypted_file = self.selected_file_path.replace(".lock", "")
            Encrypt.decrypt_file(self.selected_file_path, decrypted_file, password)
            elapsed_time = time.time() - start_time


            self.time_labels["AES Decrypt Time"].config(text=f"AES Decrypt Time: {elapsed_time:.4f} seconds")
            self.log_message(f"File decrypted in {elapsed_time:.4f} seconds: {decrypted_file}")

    def update_progress(self, value):
        """Update progress bar"""
        self.progress['value'] = value
        self.root.update_idletasks()


    def compare_aes_rsa(self):
        """Check the time log and display the bar chart"""
        # Check that all times are recorded
        times = [label.cget("text") for label in self.time_labels.values()]
        if any("Not Recorded" in time for time in times):
            messagebox.showerror("Error", "Time data is incomplete. Please perform all operations before comparison.")
            return

        # Ready time data
        aes_encrypt_time = float(self.time_labels["AES Encrypt Time"].cget("text").split(": ")[1].split()[0])
        aes_decrypt_time = float(self.time_labels["AES Decrypt Time"].cget("text").split(": ")[1].split()[0])
        rsa_encrypt_time = float(self.time_labels["RSA Encrypt Time"].cget("text").split(": ")[1].split()[0])
        rsa_decrypt_time = float(self.time_labels["RSA Decrypt Time"].cget("text").split(": ")[1].split()[0])

        # Draw a bar chart
        self.show_comparison_bar_chart(aes_encrypt_time, aes_decrypt_time, rsa_encrypt_time, rsa_decrypt_time)

    def show_comparison_bar_chart(self, aes_encrypt, aes_decrypt, rsa_encrypt, rsa_decrypt):
        """A new window pops up and displays the time comparison bar chart"""
        chart_window = tk.Toplevel(self.root)
        chart_window.title("Encryption Time Comparison Chart")
        chart_window.geometry("600x400")

        operations = ["AES Encrypt", "AES Decrypt", "RSA Encrypt", "RSA Decrypt"]
        times = [aes_encrypt, aes_decrypt, rsa_encrypt, rsa_decrypt]

        fig, ax = plt.subplots(figsize=(6, 4))
        # Draw a bar chart
        bars = ax.bar(operations, times, color=['blue', 'green', 'orange', 'red'])

        # Set the title and label
        ax.set_title("AES vs RSA Encryption Time Comparison", fontsize=14)
        ax.set_ylabel("Time (seconds)", fontsize=12)
        ax.set_xlabel("Operations", fontsize=12)

        # Adjust the Y-axis range, making sure the top is left blank to show the annotations
        max_time = max(times)
        ax.set_ylim(0, max_time + max_time * 0.2)

        for bar in bars:
            height = bar.get_height()
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                height + max_time * 0.02,
                f"{height:.4f}",
                ha='center',
                va='bottom',
                fontsize=10
            )
        # Embedded in Tkinter
        canvas = FigureCanvasTkAgg(fig, master=chart_window)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    def delete_file(self):
        """Delete selected files"""
        if not self.selected_file_path:
            messagebox.showerror("error", "Please select a file first！")
            return

        # The confirmation dialog box is displayed
        confirm = messagebox.askyesno("Confirm Delete ", f"Make sure you want to delete the file：\n{self.selected_file_path}？")
        if confirm:
            try:
                os.remove(self.selected_file_path)
                self.log_message(f"File deleted：{self.selected_file_path}")
                self.selected_file_path = None
                self.refresh_file_tree()
            except Exception as e:
                messagebox.showerror("fail to delete", f"Unable to delete file：{str(e)}")
                self.log_message(f"fail to delete：{str(e)}")
        else:
            self.log_message("Undelete operation")

    def rename_file(self):
        """Rename the selected file"""
        if not self.selected_file_path:
            messagebox.showerror("error", "Please select a file first")
            return
        new_name = simpledialog.askstring("rename file", "Please enter a new file name：")
        if new_name:
            directory = os.path.dirname(self.selected_file_path)
            new_path = os.path.join(directory, new_name)
            try:
                os.rename(self.selected_file_path, new_path)
                self.log_message(f"The file has been renamed：{new_name}")
                self.refresh_file_tree()
            except Exception as e:
                messagebox.showerror("Rename failure", f"Unable to rename file：{e}")

    def read_write_file(self):
        """Read/write selected file"""
        if not self.selected_file_path:
            messagebox.showerror("error", "Please select a file first")
            return

        # Determine whether it is a.docx file
        if self.selected_file_path.endswith(".docx"):
            try:
                document = Document(self.selected_file_path)
                content = "\n".join([paragraph.text for paragraph in document.paragraphs])
            except Exception as e:
                messagebox.showerror("Load failed", f"Unable to read file：{e}")
                return
        else:
            try:
                # Try reading a plain text file with UTF-8
                with open(self.selected_file_path, "r", encoding="utf-8") as file:
                    content = file.read()
            except UnicodeDecodeError:
                try:
                    # If UTF-8 decoding fails, try using GBK
                    with open(self.selected_file_path, "r", encoding="gbk") as file:
                        content = file.read()
                except Exception as e:
                    messagebox.showerror("Unable to read file", f"Unable to read file：{e}")
                    return

        def save_changes():
            """Save changes"""
            new_content = text_area.get("1.0", tk.END).strip()
            try:
                if self.selected_file_path.endswith(".docx"):
                    document = Document()
                    for line in new_content.splitlines():
                        document.add_paragraph(line)
                    document.save(self.selected_file_path)
                else:
                    with open(self.selected_file_path, "w", encoding="utf-8") as file:
                        file.write(new_content)
                messagebox.showinfo("save successfully", "The file was saved successfully")
                window.destroy()
            except Exception as e:
                messagebox.showerror("fail to save", f"Unable to save file：{e}")

        window = tk.Toplevel(self.root)
        window.title(f"Reading and Writing Files - {os.path.basename(self.selected_file_path)}")
        window.geometry("600x400")


        text_area = scrolledtext.ScrolledText(window, wrap=tk.WORD)
        text_area.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        text_area.insert("1.0", content)

        save_button = tk.Button(window, text="Save changes", command=save_changes)
        save_button.pack(pady=5)




if __name__ == "__main__":
    Database.initialize_database()
    root = tk.Tk()
    LoginWindow(root)
    root.mainloop()
